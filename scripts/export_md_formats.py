#!/usr/bin/env python3
"""Export teaching Markdown files to .txt and .docx siblings.

Usage:
  python scripts/export_md_formats.py notes/full-course-notes.md
  python scripts/export_md_formats.py --all
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]

DEFAULT_TARGETS = [
    "notes/full-course-notes.md",
    "notes/outline.md",
    "materials/glossary.md",
    "materials/mindmap.md",
    "materials/revision-sheet.md",
    "materials/labs/full-course-labs.md",
    "materials/qa/full-course-qa-bank.md",
    "materials/assignments/full-course-assignments.md",
    "materials/syllabus/syllabus.md",
    "materials/assignments/unit-01-assignments.md",
    "materials/assignments/unit-02-assignments.md",
    "materials/qa/unit-01-qa-bank.md",
    "materials/qa/unit-02-qa-bank.md",
]


def md_to_plain(text: str) -> str:
    lines_out: list[str] = []
    in_code = False
    for line in text.splitlines():
        if line.strip().startswith("```"):
            in_code = not in_code
            lines_out.append("")
            continue
        if in_code:
            lines_out.append(line)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            underline = {1: "=", 2: "-", 3: "~"}.get(level, "~")
            lines_out.append("")
            lines_out.append(title)
            lines_out.append(underline * min(len(title), 72))
            continue
        line = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", line)
        line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"__(.+?)__", r"\1", line)
        line = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", line)
        line = re.sub(r"`([^`]+)`", r"\1", line)
        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
                continue
            lines_out.append(" | ".join(cells))
            continue
        lines_out.append(line)
    plain = "\n".join(lines_out)
    return re.sub(r"\n{3,}", "\n\n", plain).strip() + "\n"


def add_runs_with_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            cleaned = re.sub(r"`([^`]+)`", r"\1", part)
            cleaned = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", cleaned)
            paragraph.add_run(cleaned)


def md_to_docx(md_text: str, out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_code = False
    code_buf: list[str] = []

    def flush_code() -> None:
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        code_buf = []

    for line in md_text.splitlines():
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        if not line.strip():
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            title = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(2).strip())
            doc.add_heading(title, level=level)
            continue

        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
                continue
            doc.add_paragraph(" | ".join(cells))
            continue

        if re.match(r"^[-*+]\s+", line) or re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^[-*+]\s+", "", line)
            text = re.sub(r"^\d+\.\s+", "", text)
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, text)
            continue

        p = doc.add_paragraph()
        text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", line)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        add_runs_with_bold(p, text)

    flush_code()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def export_one(md_path: Path) -> tuple[Path, Path]:
    text = md_path.read_text(encoding="utf-8")
    txt_path = md_path.with_suffix(".txt")
    docx_path = md_path.with_suffix(".docx")
    txt_path.write_text(md_to_plain(text), encoding="utf-8")
    md_to_docx(text, docx_path)
    return txt_path, docx_path


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("paths", nargs="*", type=Path, help="Markdown files to export")
    parser.add_argument("--all", action="store_true", help="Export default teaching set")
    args = parser.parse_args(argv)

    paths: list[Path] = []
    if args.all:
        paths.extend(ROOT / p for p in DEFAULT_TARGETS)
    paths.extend(p if p.is_absolute() else ROOT / p for p in args.paths)

    if not paths:
        parser.print_help()
        return 2

    for md_path in paths:
        if not md_path.exists():
            print(f"MISSING: {md_path}", file=sys.stderr)
            continue
        txt_path, docx_path = export_one(md_path)
        print(f"OK {md_path.name} -> {txt_path.name}, {docx_path.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
